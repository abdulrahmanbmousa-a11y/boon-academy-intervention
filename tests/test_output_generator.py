"""Tests for src/output_generator.py — private helper functions.

Covers _write_whatsapp_csv (OUT-03), _write_run_log (OUT-06),
_write_priority_list (OUT-01), and _write_campus_dashboards (OUT-02).
Each helper is tested independently using tmp_path for isolation.
"""
import json
from pathlib import Path

import pandas as pd
import pytest
from openpyxl import load_workbook

from src import config as cfg
from src.output_generator import (
    _write_campus_dashboards,
    _write_html_dashboard,
    _write_priority_list,
    _write_report,
    _write_run_log,
    _write_whatsapp_csv,
    write_outputs,
)


@pytest.fixture
def sample_df() -> pd.DataFrame:
    """Minimal DataFrame with 4 rows covering all risk levels.

    Columns match the full enriched DataFrame schema including LLM outputs
    and all columns required by OUTPUT_COLS_CAMPUS (15 cols).
    """
    return pd.DataFrame(
        {
            cfg.COL_STUDENT_ID: ["S001", "S002", "S003", "S004"],
            cfg.COL_STUDENT_NAME: ["Alice", "Bob", "Carol", "Dave"],
            cfg.COL_PARENT_PHONE: ["0501234567", "0502345678", "0503456789", "0504567890"],
            cfg.COL_FACILITATOR_EMAIL: [
                "fac1@boon.sa",
                "fac1@boon.sa",
                "fac2@boon.sa",
                "fac2@boon.sa",
            ],
            cfg.COL_CAMPUS_ID: ["ALPHA", "ALPHA", "BETA", "BETA"],
            cfg.COL_RISK_SCORE: [90.0, 65.0, 40.0, 15.0],
            cfg.COL_RISK_LEVEL: ["CRITICAL", "HIGH", "MEDIUM", "LOW"],
            cfg.COL_ATTENDANCE_RATE: [0.2, 0.5, 0.7, 0.9],
            cfg.COL_AVG_PRACTICE: [1.0, 3.0, 5.0, 8.0],
            cfg.COL_TREND_DIR: ["declining", "stable", "stable", "improving"],
            cfg.COL_DAYS_SINCE_NOTE: [25, 10, 5, 2],
            cfg.COL_RECOMMENDED_ACTION: [
                "Contact parent immediately",
                "Schedule check-in",
                "Monitor progress",
                "Acknowledge progress",
            ],
            cfg.COL_WHATSAPP_MESSAGE: [
                "Message for Alice",
                "Message for Bob",
                "",
                "",
            ],
            cfg.COL_GENERATED_BY: ["llm", "template", "", ""],
            cfg.COL_FACILITATOR_SUMMARY: [
                "Alice is at risk",
                "Bob needs attention",
                None,
                None,
            ],
            # Component score columns required by DISPLAY_COLS_DASHBOARD (CR-01)
            cfg.COL_ATTENDANCE_COMPONENT: [31.5, 17.5, 7.0, 1.0],
            cfg.COL_PRACTICE_COMPONENT: [27.0, 21.0, 12.0, 3.0],
            cfg.COL_TREND_COMPONENT: [20.0, 10.0, 10.0, 0.0],
            cfg.COL_NOTES_COMPONENT: [11.5, 8.0, 3.0, 0.5],
        }
    )


@pytest.fixture
def sample_run_log() -> dict:
    """Minimal run_log dict matching main.py schema (D-06)."""
    return {
        "run_timestamp": "2026-01-01T00:00:00+00:00",
        "students_processed": 10,
        "api_calls_made": 2,
        "tokens_used": {"input": 100, "output": 50},
        "errors_encountered": [],
        "fallbacks_triggered": 0,
        "data_quality_warnings": [],
    }


# ---------------------------------------------------------------------------
# _write_whatsapp_csv tests
# ---------------------------------------------------------------------------


def test_whatsapp_csv_only_critical_high(sample_df: pd.DataFrame, tmp_path: Path) -> None:
    """CSV output contains only CRITICAL and HIGH rows; MEDIUM and LOW are filtered out."""
    path = _write_whatsapp_csv(sample_df, tmp_path)
    result = pd.read_csv(path, dtype=str)
    assert set(result[cfg.COL_RISK_LEVEL].unique()) <= {"CRITICAL", "HIGH"}, (
        f"Expected only CRITICAL/HIGH in CSV, got: {result[cfg.COL_RISK_LEVEL].unique()}"
    )
    assert len(result) == 2, f"Expected 2 rows (CRITICAL+HIGH), got {len(result)}"


def test_whatsapp_csv_columns(sample_df: pd.DataFrame, tmp_path: Path) -> None:
    """CSV has exactly the required 8 columns in the specified order."""
    path = _write_whatsapp_csv(sample_df, tmp_path)
    result = pd.read_csv(path, dtype=str)
    expected_cols = [
        cfg.COL_STUDENT_ID,
        cfg.COL_STUDENT_NAME,
        cfg.COL_PARENT_PHONE,
        cfg.COL_FACILITATOR_EMAIL,
        cfg.COL_CAMPUS_ID,
        cfg.COL_RISK_LEVEL,
        cfg.COL_WHATSAPP_MESSAGE,
        cfg.COL_GENERATED_BY,
    ]
    assert list(result.columns) == expected_cols, (
        f"Column mismatch.\nExpected: {expected_cols}\nGot:      {list(result.columns)}"
    )


def test_whatsapp_csv_encoding_bom(sample_df: pd.DataFrame, tmp_path: Path) -> None:
    """File bytes start with UTF-8 BOM (b'\\xef\\xbb\\xbf') for Excel compatibility."""
    path = _write_whatsapp_csv(sample_df, tmp_path)
    raw_bytes = path.read_bytes()
    assert raw_bytes[:3] == b"\xef\xbb\xbf", (
        f"Expected UTF-8 BOM at start of file, got: {raw_bytes[:6]!r}"
    )


def test_whatsapp_csv_returns_path(sample_df: pd.DataFrame, tmp_path: Path) -> None:
    """Return value is a Path object pointing to the written file."""
    result = _write_whatsapp_csv(sample_df, tmp_path)
    assert isinstance(result, Path), f"Expected Path, got {type(result)}"
    assert result.exists(), f"Returned path does not exist: {result}"
    assert result.name == "whatsapp_messages.csv"


# ---------------------------------------------------------------------------
# _write_run_log tests
# ---------------------------------------------------------------------------


def test_run_log_keys(sample_run_log: dict, tmp_path: Path) -> None:
    """JSON output contains all 7 required top-level keys from the run_log schema."""
    path = _write_run_log(sample_run_log, tmp_path)
    data = json.loads(path.read_text(encoding="utf-8"))
    required_keys = {
        "run_timestamp",
        "students_processed",
        "api_calls_made",
        "tokens_used",
        "errors_encountered",
        "fallbacks_triggered",
        "data_quality_warnings",
    }
    missing = required_keys - set(data.keys())
    assert not missing, f"Missing keys in run_log.json: {missing}"


def test_run_log_returns_path(sample_run_log: dict, tmp_path: Path) -> None:
    """Return value is a Path object pointing to the written JSON file."""
    result = _write_run_log(sample_run_log, tmp_path)
    assert isinstance(result, Path), f"Expected Path, got {type(result)}"
    assert result.exists(), f"Returned path does not exist: {result}"
    assert result.name == "run_log.json"


def test_run_log_default_str_handles_non_serializable(tmp_path: Path) -> None:
    """run_log containing a datetime object serializes without raising TypeError."""
    import datetime

    run_log_with_datetime = {
        "run_timestamp": datetime.datetime(2026, 1, 1, tzinfo=datetime.timezone.utc),
        "students_processed": 5,
        "api_calls_made": 1,
        "tokens_used": {"input": 50, "output": 25},
        "errors_encountered": [],
        "fallbacks_triggered": 0,
        "data_quality_warnings": [],
    }
    # Must not raise TypeError — default=str handles datetime objects
    path = _write_run_log(run_log_with_datetime, tmp_path)
    data = json.loads(path.read_text(encoding="utf-8"))
    # datetime was serialized as a string — confirm it is readable
    assert isinstance(data["run_timestamp"], str), (
        f"Expected run_timestamp as str, got {type(data['run_timestamp'])}"
    )


# ---------------------------------------------------------------------------
# _write_priority_list fixtures and tests (OUT-01)
# ---------------------------------------------------------------------------


@pytest.fixture
def priority_list_path(sample_df: pd.DataFrame, tmp_path: Path) -> Path:
    """Write priority list to tmp_path and return the path for round-trip assertions."""
    return _write_priority_list(sample_df, tmp_path)


def test_priority_list_file_exists(priority_list_path: Path) -> None:
    """_write_priority_list writes intervention_priority_list.xlsx to output_dir."""
    assert priority_list_path.exists(), f"File not found: {priority_list_path}"
    assert priority_list_path.name == "intervention_priority_list.xlsx"


def test_priority_list_header_color(priority_list_path: Path) -> None:
    """Header cell A1 has navy fill after save+reload (FF1F4E79)."""
    wb = load_workbook(priority_list_path)
    ws = wb.active
    assert ws["A1"].fill.fgColor.rgb == cfg.COLOR_HEADER, (
        f"Expected header fill {cfg.COLOR_HEADER}, got {ws['A1'].fill.fgColor.rgb}"
    )


def test_priority_list_header_font(priority_list_path: Path) -> None:
    """Header cell A1 has bold white font after save+reload."""
    wb = load_workbook(priority_list_path)
    ws = wb.active
    assert ws["A1"].font.bold is True, "Expected A1 font.bold to be True"
    assert ws["A1"].font.color.rgb == cfg.FONT_WHITE, (
        f"Expected header font color {cfg.FONT_WHITE}, got {ws['A1'].font.color.rgb}"
    )


def test_priority_list_freeze_panes(priority_list_path: Path) -> None:
    """ws.freeze_panes == 'A2' after save+reload — header row is frozen."""
    wb = load_workbook(priority_list_path)
    ws = wb.active
    assert ws.freeze_panes == "A2", (
        f"Expected freeze_panes='A2', got {ws.freeze_panes!r}"
    )


def test_priority_list_critical_row_color(priority_list_path: Path) -> None:
    """First data row (row 2) has CRITICAL fill color after save+reload."""
    wb = load_workbook(priority_list_path)
    ws = wb.active
    # sample_df has CRITICAL (score=90) as first row — sorted desc it must be row 2
    assert ws["A2"].fill.fgColor.rgb == cfg.COLOR_CRITICAL, (
        f"Expected CRITICAL fill {cfg.COLOR_CRITICAL} at A2, got {ws['A2'].fill.fgColor.rgb}"
    )


def test_priority_list_sorted_desc(priority_list_path: Path) -> None:
    """Rows are sorted by risk_score descending — row 2 has higher score than row 3."""
    wb = load_workbook(priority_list_path)
    ws = wb.active
    # risk_score is column index 6 in OUTPUT_COLS_PRIORITY (1-based)
    risk_score_col = cfg.OUTPUT_COLS_PRIORITY.index(cfg.COL_RISK_SCORE) + 1
    score_row2 = ws.cell(row=2, column=risk_score_col).value
    score_row3 = ws.cell(row=3, column=risk_score_col).value
    assert score_row2 > score_row3, (
        f"Expected row 2 score ({score_row2}) > row 3 score ({score_row3}) — must be sorted desc"
    )


def test_priority_list_rank_column(priority_list_path: Path) -> None:
    """First data row has rank=1 in the rank column (A2)."""
    wb = load_workbook(priority_list_path)
    ws = wb.active
    assert ws["A2"].value == 1, (
        f"Expected rank=1 at A2 (first data row), got {ws['A2'].value}"
    )


def test_priority_list_column_count(priority_list_path: Path) -> None:
    """Worksheet has exactly 12 columns (OUTPUT_COLS_PRIORITY length)."""
    wb = load_workbook(priority_list_path)
    ws = wb.active
    assert ws.max_column == 12, (
        f"Expected 12 columns (OUTPUT_COLS_PRIORITY), got {ws.max_column}"
    )


# ---------------------------------------------------------------------------
# _write_campus_dashboards fixtures and tests (OUT-02)
# ---------------------------------------------------------------------------


@pytest.fixture
def multi_campus_df() -> pd.DataFrame:
    """DataFrame with 2 campuses plus a NaN-campus row for dropna=True validation.

    ALPHA: 1 CRITICAL + 1 MEDIUM
    BETA:  1 HIGH + 1 LOW
    Plus one row with campus_id=NaN (must be excluded from output files).
    All 15 OUTPUT_COLS_CAMPUS columns are present. MEDIUM and LOW rows
    have None in the 3 LLM columns per D-06.
    """
    import math

    return pd.DataFrame(
        {
            cfg.COL_STUDENT_ID: ["S001", "S002", "S003", "S004", "S005"],
            cfg.COL_STUDENT_NAME: ["Alice", "Carol", "Bob", "Dave", "Eve"],
            cfg.COL_PARENT_PHONE: [
                "0501111111",
                "0503333333",
                "0502222222",
                "0504444444",
                "0505555555",
            ],
            cfg.COL_FACILITATOR_EMAIL: [
                "fac@alpha.sa",
                "fac@alpha.sa",
                "fac@beta.sa",
                "fac@beta.sa",
                "fac@unknown.sa",
            ],
            cfg.COL_CAMPUS_ID: ["ALPHA", "ALPHA", "BETA", "BETA", float("nan")],
            cfg.COL_RISK_SCORE: [90.0, 35.0, 70.0, 15.0, 50.0],
            cfg.COL_RISK_LEVEL: ["CRITICAL", "MEDIUM", "HIGH", "LOW", "HIGH"],
            cfg.COL_ATTENDANCE_RATE: [0.2, 0.7, 0.5, 0.9, 0.6],
            cfg.COL_AVG_PRACTICE: [1.0, 5.0, 3.0, 8.0, 4.0],
            cfg.COL_TREND_DIR: ["declining", "stable", "stable", "improving", "stable"],
            cfg.COL_DAYS_SINCE_NOTE: [25, 5, 10, 2, 7],
            cfg.COL_RECOMMENDED_ACTION: [
                "Contact parent immediately",
                "Monitor progress",
                "Schedule check-in",
                "Acknowledge progress",
                "Schedule check-in",
            ],
            cfg.COL_FACILITATOR_SUMMARY: [
                "Alice is at critical risk",
                None,
                "Bob needs attention",
                None,
                "Eve needs attention",
            ],
            cfg.COL_WHATSAPP_MESSAGE: [
                "Message for Alice",
                None,
                "Message for Bob",
                None,
                "Message for Eve",
            ],
            cfg.COL_GENERATED_BY: ["llm", None, "template", None, "llm"],
        }
    )


@pytest.fixture
def campus_dashboard_paths(
    multi_campus_df: pd.DataFrame, tmp_path: Path
) -> dict[str, Path]:
    """Write campus dashboards to tmp_path and return the result dict."""
    return _write_campus_dashboards(multi_campus_df, tmp_path)


def test_campus_dashboard_files_created(
    campus_dashboard_paths: dict[str, Path], tmp_path: Path
) -> None:
    """_write_campus_dashboards returns dict with keys for ALPHA and BETA; files exist."""
    assert "campus_ALPHA" in campus_dashboard_paths, (
        f"Missing key 'campus_ALPHA' in {list(campus_dashboard_paths.keys())}"
    )
    assert "campus_BETA" in campus_dashboard_paths, (
        f"Missing key 'campus_BETA' in {list(campus_dashboard_paths.keys())}"
    )
    assert campus_dashboard_paths["campus_ALPHA"].exists()
    assert campus_dashboard_paths["campus_BETA"].exists()


def test_campus_dashboard_header_row(campus_dashboard_paths: dict[str, Path]) -> None:
    """Header cell A1 has navy fill and bold font after save+reload."""
    wb = load_workbook(campus_dashboard_paths["campus_ALPHA"])
    ws = wb.active
    assert ws["A1"].fill.fgColor.rgb == cfg.COLOR_HEADER, (
        f"Expected header fill {cfg.COLOR_HEADER}, got {ws['A1'].fill.fgColor.rgb}"
    )
    assert ws["A1"].font.bold is True, "Expected A1 font.bold to be True"


def test_campus_dashboard_freeze_panes(campus_dashboard_paths: dict[str, Path]) -> None:
    """ws.freeze_panes == 'A2' after save+reload."""
    wb = load_workbook(campus_dashboard_paths["campus_ALPHA"])
    ws = wb.active
    assert ws.freeze_panes == "A2", (
        f"Expected freeze_panes='A2', got {ws.freeze_panes!r}"
    )


def test_campus_dashboard_column_count(campus_dashboard_paths: dict[str, Path]) -> None:
    """Worksheet has exactly 15 columns (OUTPUT_COLS_CAMPUS length)."""
    wb = load_workbook(campus_dashboard_paths["campus_ALPHA"])
    ws = wb.active
    assert ws.max_column == 15, (
        f"Expected 15 columns (OUTPUT_COLS_CAMPUS), got {ws.max_column}"
    )


def test_campus_dashboard_summary_row(campus_dashboard_paths: dict[str, Path]) -> None:
    """Row 2 is the summary row — cell A2 contains 'Summary' text."""
    wb = load_workbook(campus_dashboard_paths["campus_ALPHA"])
    ws = wb.active
    assert ws["A2"].value == "Summary", (
        f"Expected 'Summary' in A2 (summary row), got {ws['A2'].value!r}"
    )


def test_campus_dashboard_data_starts_row3(
    campus_dashboard_paths: dict[str, Path],
) -> None:
    """Data rows start at row 3 — A3 contains rank 1 (first student by risk_score)."""
    wb = load_workbook(campus_dashboard_paths["campus_ALPHA"])
    ws = wb.active
    assert ws["A3"].value == 1, (
        f"Expected rank=1 at A3 (first data row), got {ws['A3'].value}"
    )


def test_campus_dashboard_critical_row_color(
    campus_dashboard_paths: dict[str, Path],
) -> None:
    """CRITICAL student data row (row 3+) has CRITICAL fill color after save+reload."""
    wb = load_workbook(campus_dashboard_paths["campus_ALPHA"])
    ws = wb.active
    # ALPHA has CRITICAL student as highest-scored — should be row 3 (first data row)
    assert ws["A3"].fill.fgColor.rgb == cfg.COLOR_CRITICAL, (
        f"Expected CRITICAL fill {cfg.COLOR_CRITICAL} at A3, got {ws['A3'].fill.fgColor.rgb}"
    )


def test_campus_dashboard_medium_llm_cells_empty(
    campus_dashboard_paths: dict[str, Path],
) -> None:
    """MEDIUM student rows have None (empty) in LLM columns 13, 14, 15 (D-06)."""
    wb = load_workbook(campus_dashboard_paths["campus_ALPHA"])
    ws = wb.active
    risk_col_idx = cfg.OUTPUT_COLS_CAMPUS.index(cfg.COL_RISK_LEVEL) + 1
    # Find the MEDIUM row (row 3 onward — skip header row 1 and summary row 2)
    medium_row = None
    for row in ws.iter_rows(min_row=3):
        if row[risk_col_idx - 1].value == "MEDIUM":
            medium_row = row
            break
    assert medium_row is not None, "Could not find MEDIUM student row in ALPHA dashboard"
    # Columns 13 (facilitator_summary), 14 (whatsapp_message), 15 (generated_by) must be None
    assert medium_row[12].value is None, (
        f"Expected None in col 13 (facilitator_summary) for MEDIUM row, got {medium_row[12].value!r}"
    )
    assert medium_row[13].value is None, (
        f"Expected None in col 14 (whatsapp_message) for MEDIUM row, got {medium_row[13].value!r}"
    )
    assert medium_row[14].value is None, (
        f"Expected None in col 15 (generated_by) for MEDIUM row, got {medium_row[14].value!r}"
    )


def test_campus_dashboard_excludes_nan_campus(
    campus_dashboard_paths: dict[str, Path], tmp_path: Path
) -> None:
    """No 'facilitator_dashboard_nan.xlsx' is created — dropna=True in groupby."""
    nan_file = tmp_path / "facilitator_dashboard_nan.xlsx"
    assert not nan_file.exists(), (
        "facilitator_dashboard_nan.xlsx must not be created (dropna=True in groupby)"
    )
    assert "campus_nan" not in campus_dashboard_paths, (
        "Key 'campus_nan' must not appear in results dict"
    )


# ---------------------------------------------------------------------------
# write_outputs integration tests (OUT-01, OUT-02, OUT-03, OUT-06)
# ---------------------------------------------------------------------------


@pytest.fixture
def full_sample_df() -> pd.DataFrame:
    """Full enriched DataFrame with 2 campuses (ALPHA, BETA), covering all risk levels.

    Matches the real pipeline schema: student_id and parent_phone as str dtype.
    Includes at least 1 CRITICAL and 1 HIGH student per requirement.
    All 15 OUTPUT_COLS_CAMPUS columns are present.
    """
    return pd.DataFrame(
        {
            cfg.COL_STUDENT_ID: pd.array(
                ["S001", "S002", "S003", "S004", "S005", "S006"], dtype="string"
            ),
            cfg.COL_STUDENT_NAME: ["Alice", "Bob", "Carol", "Dave", "Eve", "Frank"],
            cfg.COL_PARENT_PHONE: pd.array(
                [
                    "0501111111",
                    "0502222222",
                    "0503333333",
                    "0504444444",
                    "0505555555",
                    "0506666666",
                ],
                dtype="string",
            ),
            cfg.COL_FACILITATOR_EMAIL: [
                "fac@alpha.sa",
                "fac@alpha.sa",
                "fac@alpha.sa",
                "fac@beta.sa",
                "fac@beta.sa",
                "fac@beta.sa",
            ],
            cfg.COL_CAMPUS_ID: ["ALPHA", "ALPHA", "ALPHA", "BETA", "BETA", "BETA"],
            cfg.COL_RISK_SCORE: [92.0, 68.0, 38.0, 85.0, 55.0, 12.0],
            cfg.COL_RISK_LEVEL: ["CRITICAL", "HIGH", "MEDIUM", "CRITICAL", "HIGH", "LOW"],
            cfg.COL_ATTENDANCE_RATE: [0.1, 0.45, 0.7, 0.15, 0.55, 0.95],
            cfg.COL_AVG_PRACTICE: [0.5, 2.5, 5.0, 1.0, 4.0, 9.0],
            cfg.COL_TREND_DIR: [
                "declining",
                "stable",
                "stable",
                "declining",
                "stable",
                "improving",
            ],
            cfg.COL_DAYS_SINCE_NOTE: [30, 12, 4, 28, 8, 1],
            cfg.COL_RECOMMENDED_ACTION: [
                "Contact parent immediately",
                "Schedule check-in",
                "Monitor progress",
                "Contact parent immediately",
                "Schedule check-in",
                "Acknowledge progress",
            ],
            cfg.COL_FACILITATOR_SUMMARY: [
                "Alice is at critical risk",
                "Bob needs attention",
                None,
                "Dave is at critical risk",
                "Eve needs attention",
                None,
            ],
            cfg.COL_WHATSAPP_MESSAGE: [
                "Message for Alice",
                "Message for Bob",
                None,
                "Message for Dave",
                "Message for Eve",
                None,
            ],
            cfg.COL_GENERATED_BY: ["llm", "template", None, "llm", "template", None],
            # Component score columns required by DISPLAY_COLS_DASHBOARD (CR-01)
            cfg.COL_ATTENDANCE_COMPONENT: [32.2, 15.8, 5.5, 30.1, 18.0, 1.5],
            cfg.COL_PRACTICE_COMPONENT: [28.5, 20.0, 10.0, 25.0, 22.0, 3.0],
            cfg.COL_TREND_COMPONENT: [20.0, 10.0, 10.0, 20.0, 10.0, 0.0],
            cfg.COL_NOTES_COMPONENT: [11.3, 8.0, 3.0, 10.5, 5.0, 0.5],
        }
    )


@pytest.fixture
def sample_run_log_full() -> dict:
    """Full run_log dict matching the 7-key D-06 schema from main.py."""
    return {
        "run_timestamp": "2026-05-23T16:00:00+00:00",
        "students_processed": 6,
        "api_calls_made": 2,
        "tokens_used": {"input": 500, "output": 200},
        "errors_encountered": [],
        "fallbacks_triggered": 0,
        "data_quality_warnings": [],
    }


def test_write_outputs_returns_all_keys(
    full_sample_df: pd.DataFrame,
    sample_run_log_full: dict,
    tmp_path: Path,
) -> None:
    """write_outputs returns dict with 'priority_list', campus_* keys, 'whatsapp', 'run_log'."""
    result = write_outputs(full_sample_df, tmp_path, sample_run_log_full)
    assert "priority_list" in result, (
        f"Missing key 'priority_list' in result: {list(result.keys())}"
    )
    campus_keys = [k for k in result if k.startswith("campus_")]
    assert len(campus_keys) >= 1, (
        f"Expected at least one 'campus_*' key, got: {list(result.keys())}"
    )
    assert "whatsapp" in result, (
        f"Missing key 'whatsapp' in result: {list(result.keys())}"
    )
    assert "run_log" in result, (
        f"Missing key 'run_log' in result: {list(result.keys())}"
    )
    assert "dashboard" in result, (
        f"Missing key 'dashboard' in result: {list(result.keys())}"
    )
    assert "report" in result, (
        f"Missing key 'report' in result: {list(result.keys())}"
    )


def test_write_outputs_all_paths_exist(
    full_sample_df: pd.DataFrame,
    sample_run_log_full: dict,
    tmp_path: Path,
) -> None:
    """Every Path value in the write_outputs return dict points to a file that exists."""
    result = write_outputs(full_sample_df, tmp_path, sample_run_log_full)
    for key, path in result.items():
        assert isinstance(path, Path), f"Expected Path for key {key!r}, got {type(path)}"
        assert path.exists(), f"File does not exist for key {key!r}: {path}"


def test_all_6_output_files_exist(
    minimal_enriched_df: pd.DataFrame,
    tmp_path: Path,
) -> None:
    """TEST-04 / D-06: write_outputs() produces all 6 output files in tmp_path.

    Uses the shared minimal_enriched_df fixture from conftest.py (5 rows, 2 campuses)
    to call write_outputs() end-to-end. Every Path value in the returned dict must
    exist on disk. All 5 non-campus fixed keys (priority_list, whatsapp, run_log,
    dashboard, report) must be present in the result.
    """
    run_log = {
        "run_timestamp": "2026-05-24T12:00:00+00:00",
        "students_processed": 5,
        "api_calls_made": 2,
        "tokens_used": {"input": 150, "output": 80},
        "errors_encountered": [],
        "fallbacks_triggered": 0,
        "data_quality_warnings": [],
    }
    result = write_outputs(minimal_enriched_df, tmp_path, run_log)
    # Every value must be a Path that exists on disk
    for key, path in result.items():
        assert isinstance(path, Path), (
            f"TEST-04: expected Path for key {key!r}, got {type(path)}"
        )
        assert path.exists(), (
            f"TEST-04: file does not exist for key {key!r}: {path}"
        )
    # Assert the 5 mandatory non-campus keys are present
    for required_key in ("priority_list", "whatsapp", "run_log", "dashboard", "report"):
        assert required_key in result, (
            f"TEST-04 / D-06: missing required key {required_key!r} in write_outputs result"
        )


def test_write_outputs_creates_output_dir(
    full_sample_df: pd.DataFrame,
    sample_run_log_full: dict,
    tmp_path: Path,
) -> None:
    """write_outputs creates a non-existent output directory without raising an error."""
    nested_dir = tmp_path / "deep" / "nested" / "outputs"
    assert not nested_dir.exists(), "Pre-condition: directory must not exist before call"
    # Must not raise — output_dir.mkdir(parents=True, exist_ok=True) handles it
    result = write_outputs(full_sample_df, nested_dir, sample_run_log_full)
    assert nested_dir.exists(), f"Expected output_dir to be created: {nested_dir}"
    assert len(result) > 0, "Expected non-empty result dict after write_outputs call"


# ---------------------------------------------------------------------------
# _write_html_dashboard tests (OUT-05)
# ---------------------------------------------------------------------------


@pytest.fixture
def html_dashboard_path(sample_df: pd.DataFrame, tmp_path: Path) -> Path:
    """Write HTML dashboard to tmp_path and return the path for round-trip assertions."""
    return _write_html_dashboard(sample_df, tmp_path)


def test_html_dashboard_returns_path(
    sample_df: pd.DataFrame, tmp_path: Path
) -> None:
    """_write_html_dashboard returns a Path pointing to facilitator_dashboard.html."""
    result = _write_html_dashboard(sample_df, tmp_path)
    assert isinstance(result, Path), f"Expected Path, got {type(result)}"
    assert result.exists(), f"Returned path does not exist: {result}"
    assert result.name == "facilitator_dashboard.html"


def test_html_dashboard_contains_student_data(html_dashboard_path: Path) -> None:
    """Rendered HTML embeds the JS const studentsData with at least one student ID."""
    content = html_dashboard_path.read_text(encoding="utf-8")
    assert "studentsData" in content, "Expected JS const 'studentsData' in HTML"
    assert "S001" in content, "Expected student_id S001 in embedded JSON"


def test_html_dashboard_contains_campus_ids(html_dashboard_path: Path) -> None:
    """Campus IDs from sample_df appear in the rendered HTML (campus filter options)."""
    content = html_dashboard_path.read_text(encoding="utf-8")
    assert "ALPHA" in content, "Expected campus_id 'ALPHA' in rendered HTML"
    assert "BETA" in content, "Expected campus_id 'BETA' in rendered HTML"


def test_html_dashboard_no_external_urls(html_dashboard_path: Path) -> None:
    """HTML file contains no external http:// or https:// URLs — fully self-contained."""
    content = html_dashboard_path.read_text(encoding="utf-8")
    assert "https://" not in content, "Unexpected https:// URL found — file must be self-contained"
    assert "http://" not in content, "Unexpected http:// URL found — file must be self-contained"


def test_html_dashboard_escape_script_tag(
    sample_df: pd.DataFrame, tmp_path: Path
) -> None:
    """json.dumps().replace('</','<\\/') prevents </script> injection in the JSON data block.

    Builds a one-row DataFrame with a whatsapp_message containing a literal
    </script> tag, renders the dashboard, then asserts the embedded JSON data
    block does not contain an unescaped </script> sequence.
    """
    # Build minimal one-row DataFrame with injection payload
    injection_df = pd.DataFrame(
        {
            cfg.COL_STUDENT_ID: ["INJ01"],
            cfg.COL_STUDENT_NAME: ["Test Student"],
            cfg.COL_CAMPUS_ID: ["CAMPUS_X"],
            cfg.COL_RISK_SCORE: [80.0],
            cfg.COL_RISK_LEVEL: ["CRITICAL"],
            cfg.COL_ATTENDANCE_RATE: [0.3],
            cfg.COL_AVG_PRACTICE: [2.0],
            cfg.COL_TREND_DIR: ["declining"],
            cfg.COL_DAYS_SINCE_NOTE: [15],
            cfg.COL_FACILITATOR_SUMMARY: ["Summary text"],
            cfg.COL_WHATSAPP_MESSAGE: ["</script><script>alert(1)</script>"],
            cfg.COL_GENERATED_BY: ["llm"],
            # Component columns required by DISPLAY_COLS_DASHBOARD (CR-01)
            cfg.COL_ATTENDANCE_COMPONENT: [28.0],
            cfg.COL_PRACTICE_COMPONENT: [24.0],
            cfg.COL_TREND_COMPONENT: [20.0],
            cfg.COL_NOTES_COMPONENT: [8.0],
        }
    )
    path = _write_html_dashboard(injection_df, tmp_path)
    content = path.read_text(encoding="utf-8")

    # Find the JSON data block that starts with the const declaration
    json_start = content.find("const studentsData =")
    assert json_start != -1, "Could not find 'const studentsData =' in rendered HTML"

    # Extract from the const declaration to the closing semicolon of the array
    # The JSON array ends with "];", look for the next ";" after the opening bracket
    json_block_start = content.find("[", json_start)
    json_block_end = content.find("];", json_block_start)
    assert json_block_start != -1 and json_block_end != -1, (
        "Could not delimit the studentsData JSON array in rendered HTML"
    )
    json_block = content[json_block_start : json_block_end + 2]

    # The literal string "</script>" must NOT appear in the JSON data block
    assert "</script>" not in json_block, (
        "Unescaped </script> found in studentsData JSON block — "
        "json.dumps().replace('</','<\\/') injection guard is broken"
    )


# ---------------------------------------------------------------------------
# _write_report tests (OUT-04)
# ---------------------------------------------------------------------------


@pytest.fixture
def report_path(sample_df: pd.DataFrame, sample_run_log: dict, tmp_path: Path) -> Path:
    """Write Word intervention report to tmp_path and return the path for assertions."""
    return _write_report(sample_df, sample_run_log, tmp_path)


def test_report_returns_path(report_path: Path) -> None:
    """_write_report returns a Path pointing to intervention_report.docx."""
    assert isinstance(report_path, Path), f"Expected Path, got {type(report_path)}"
    assert report_path.exists(), f"Returned path does not exist: {report_path}"
    assert report_path.name == "intervention_report.docx"


def test_report_is_valid_docx(report_path: Path) -> None:
    """intervention_report.docx is a valid Word document with at least one paragraph."""
    from docx import Document

    doc = Document(str(report_path))
    assert len(doc.paragraphs) > 0, "Expected non-empty Word document"


def test_report_contains_cover_heading(report_path: Path) -> None:
    """Cover page heading contains 'Intervention Report'.

    python-docx maps level=0 to the 'Title' style (not 'Heading 1'), so we
    check both 'Heading' and 'Title' style paragraphs.
    """
    from docx import Document

    doc = Document(str(report_path))
    # level=0 -> "Title" style; level=1/2 -> "Heading N" style
    heading_texts = [
        p.text
        for p in doc.paragraphs
        if p.style.name.startswith("Heading") or p.style.name == "Title"
    ]
    assert any("Intervention Report" in h for h in heading_texts), (
        f"Cover heading not found. Heading/Title paragraphs: {heading_texts}"
    )


def test_report_contains_executive_summary(report_path: Path) -> None:
    """Document contains an 'Executive Summary' section heading."""
    from docx import Document

    doc = Document(str(report_path))
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert any("Executive Summary" in h for h in headings), (
        f"'Executive Summary' heading not found. Headings: {headings}"
    )


def test_report_contains_tables(report_path: Path) -> None:
    """Document contains at least 3 tables (risk breakdown, top-10, campus summary)."""
    from docx import Document

    doc = Document(str(report_path))
    assert len(doc.tables) >= 3, (
        f"Expected at least 3 tables (risk breakdown, top-10, campus summary), "
        f"got {len(doc.tables)}"
    )


def test_report_data_quality_no_warnings(report_path: Path) -> None:
    """When data_quality_warnings is empty, document contains 'No data quality issues'."""
    from docx import Document

    doc = Document(str(report_path))
    all_text = " ".join(p.text for p in doc.paragraphs)
    assert "No data quality issues" in all_text, (
        "Expected 'No data quality issues' paragraph when warnings list is empty"
    )


def test_write_outputs_html_contains_embedded_json(
    full_sample_df: pd.DataFrame,
    sample_run_log_full: dict,
    tmp_path: Path,
) -> None:
    """write_outputs integration: facilitator_dashboard.html contains const studentsData."""
    result = write_outputs(full_sample_df, tmp_path, sample_run_log_full)
    html_path = result["dashboard"]
    content = html_path.read_text(encoding="utf-8")
    assert "const studentsData" in content, (
        f"Expected 'const studentsData' in {html_path.name}, got first 200 chars: {content[:200]}"
    )
    assert "S001" in content, (
        "Expected student_id S001 to appear in embedded JSON"
    )
