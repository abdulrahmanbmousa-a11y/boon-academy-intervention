"""Documentation generation module for boon-academy-intervention.

Phase 6 implements write_docs() and its nine private helper stubs for writing
all documentation output files from the enriched DataFrame and run_log.

D-01 (Signature): write_docs(df, run_log, docs_dir) — produces 9 documentation
files: analysis.md (Markdown at project root), analysis.docx, and seven static
technical .docx files (architecture, security, engineering_decisions, data_handling,
scalability, system_design, alternatives).

D-04 / D-05 (Content loading): Static content for the 7 .docx files is loaded
once at module import from src/templates/docs_content.yaml via yaml.safe_load().
Identical pattern to llm_templates.yaml in src/llm_engine.py.

D-08 (Module discipline): logging.getLogger(__name__), type hints on all functions,
docstrings on public functions, zero print() statements.

D-09 (Helper list): Nine private _write_* helpers — one per output file — each
returning a Path.  Implementations are Wave 2; this module provides the skeleton
so Wave 2 plans can implement helpers independently.

Security:
- T-06-02: No ANTHROPIC_API_KEY or student PII in any logger.* call — stubs log path only
- T-06-03: yaml.safe_load() (not yaml.load()) — tamper-resistant YAML loading
"""
import logging
from pathlib import Path

import pandas as pd
import yaml
from docx import Document

from src import config as cfg

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Module-level YAML load — loaded once at import time (D-04, D-05)
# ---------------------------------------------------------------------------

_CONTENT_PATH = Path(__file__).parent / "templates" / "docs_content.yaml"
with _CONTENT_PATH.open(encoding="utf-8") as _fh:
    _DOCS_CONTENT: dict = yaml.safe_load(_fh)


# ---------------------------------------------------------------------------
# Shared rendering helper
# ---------------------------------------------------------------------------

def _render_doc_from_content(content_dict: dict, docs_dir: Path, filename: str) -> Path:
    """Render a YAML content dict into a .docx file using python-docx.

    Renders a single documentation file from a parsed YAML content block.
    Each content_dict must have a top-level 'title' (str) and 'sections' (list).
    Each section has 'heading' (str) and 'body' (list of {type, text} dicts).
    Supported body types: 'paragraph' and 'bullet'.

    D-10: Uses only built-in python-docx styles — add_heading(level=0/1),
    add_paragraph(), paragraph.style. No OxmlElement, no custom styles.

    Args:
        content_dict: Parsed YAML dict with 'title' and 'sections' keys.
        docs_dir: Directory to write the .docx file into.
        filename: Output filename (e.g. 'architecture.docx').

    Returns:
        Path to the written .docx file.
    """
    doc = Document()

    title = content_dict.get("title", filename.replace(".docx", "").replace("_", " ").title())
    doc.add_heading(title, level=0)

    for section in content_dict.get("sections", []):
        heading = section.get("heading", "")
        if heading:
            doc.add_heading(heading, level=1)

        for item in section.get("body", []):
            item_type = item.get("type", "paragraph")
            text = item.get("text", "")
            if item_type == "bullet":
                para = doc.add_paragraph(text, style="List Bullet")
            else:
                doc.add_paragraph(text)

    path = docs_dir / filename
    # str() required for python-docx 1.1.2 on Windows (CLAUDE.md critical pitfall)
    doc.save(str(path))
    logger.info("Wrote doc: %s", path)
    return path


# ---------------------------------------------------------------------------
# Public orchestrator
# ---------------------------------------------------------------------------

def write_docs(df: pd.DataFrame, run_log: dict, docs_dir: Path) -> dict[str, Path]:
    """Write all 9 documentation files and return a dict of named paths.

    Creates docs_dir (including parents) if it does not exist, then delegates
    to nine private helpers — one per output file — and returns a unified dict
    mapping semantic keys to resolved Path objects (D-01, D-03).

    D-09 helper list:
      _write_analysis_md  — analysis.md at project root (Markdown, live run_log numbers)
      _write_analysis_docx — docs/analysis.docx (Word version of analysis.md)
      _write_architecture  — docs/architecture.docx
      _write_security      — docs/security.docx
      _write_engineering_decisions — docs/engineering_decisions.docx
      _write_data_handling — docs/data_handling.docx
      _write_scalability   — docs/scalability.docx
      _write_system_design — docs/system_design.docx
      _write_alternatives  — docs/alternatives.docx

    Args:
        df: Fully enriched one-row-per-student DataFrame from enrich_with_llm().
            Used by _write_analysis_md to compute risk-level distribution.
        run_log: In-memory pipeline run metadata dict with keys: run_timestamp,
            students_processed, api_calls_made, tokens_used, errors_encountered,
            fallbacks_triggered, data_quality_warnings.
        docs_dir: Directory to write .docx files into (cfg.DOCS_DIR).

    Returns:
        Dict mapping output file keys to their resolved Path objects.
        Keys: analysis_md, analysis_docx, architecture, security,
              engineering_decisions, data_handling, scalability,
              system_design, alternatives.
    """
    docs_dir.mkdir(parents=True, exist_ok=True)  # D-12

    paths: dict[str, Path] = {}

    paths["analysis_md"] = _write_analysis_md(df, run_log, docs_dir)
    paths["analysis_docx"] = _write_analysis_docx(run_log, docs_dir, _DOCS_CONTENT)
    paths["architecture"] = _write_architecture(docs_dir, _DOCS_CONTENT.get("architecture", {}))
    paths["security"] = _write_security(docs_dir, _DOCS_CONTENT.get("security", {}))
    paths["engineering_decisions"] = _write_engineering_decisions(
        docs_dir, _DOCS_CONTENT.get("engineering_decisions", {})
    )
    paths["data_handling"] = _write_data_handling(
        docs_dir, _DOCS_CONTENT.get("data_handling", {})
    )
    paths["scalability"] = _write_scalability(
        docs_dir, _DOCS_CONTENT.get("scalability", {})
    )
    paths["system_design"] = _write_system_design(
        docs_dir, _DOCS_CONTENT.get("system_design", {})
    )
    paths["alternatives"] = _write_alternatives(
        docs_dir, _DOCS_CONTENT.get("alternatives", {})
    )

    logger.info("docs written: %s", list(paths.keys()))
    return paths


# ---------------------------------------------------------------------------
# Private helpers — stubs (Wave 2 will implement each)
# ---------------------------------------------------------------------------

def _write_analysis_md(df: pd.DataFrame, run_log: dict, docs_dir: Path) -> Path:
    """Write analysis.md at project root as a 5-section Markdown memo.

    Produces a plain Markdown memo (under 600 words) with 5 sections:
    Diagnosis, What You Found, What You Built, What You Cut, What Next.
    Embeds live run_log numbers (students_processed, api_calls_made, tokens_used)
    and the runtime risk distribution from df[COL_RISK_LEVEL].value_counts().

    analysis.md is written at the project root (docs_dir.parent / "analysis.md"),
    not inside docs_dir, because it is the primary deliverable memo (D-11).

    Security: No student names, parent phones, or API key are embedded — only
    aggregate counts from run_log and df (T-06-04).

    Args:
        df: Enriched DataFrame — used for risk-level distribution counts (D-07).
        run_log: Pipeline run metadata dict with live numbers. Keys accessed via
            .get() with defaults to tolerate partial run_logs (T-06-05).
        docs_dir: docs/ directory — parent is the project root for analysis.md.

    Returns:
        Path to analysis.md at project root.
    """
    path = docs_dir.parent / "analysis.md"

    # Risk distribution from DataFrame at runtime
    risk_dist = df[cfg.COL_RISK_LEVEL].value_counts().to_dict()
    n_critical = risk_dist.get("CRITICAL", 0)
    n_high = risk_dist.get("HIGH", 0)
    n_medium = risk_dist.get("MEDIUM", 0)
    n_low = risk_dist.get("LOW", 0)

    # Token counts
    tokens_used = run_log.get("tokens_used", {})
    tokens_in = tokens_used.get("input", 0)
    tokens_out = tokens_used.get("output", 0)
    tokens_total = tokens_in + tokens_out

    # Pipeline stats
    students_processed = run_log.get("students_processed", 0)
    api_calls = run_log.get("api_calls_made", 0)
    fallbacks = run_log.get("fallbacks_triggered", 0)

    # Duplicate ID count from data_quality_warnings
    dup_count = sum(
        1
        for w in run_log.get("data_quality_warnings", [])
        if isinstance(w, dict) and w.get("type") == "duplicate_id"
    )

    content = f"""# Analysis: boon-academy-intervention

## Diagnosis

Boon Academy runs 20 campuses serving roughly 300 students. Facilitator
intervention rate is currently ~30% against an 80%+ target. The gap exists
because facilitators lack a fast, prioritised view of which students to contact
and what to say. This pipeline closes that gap by scoring every student daily,
ranking them by risk, and generating draft WhatsApp messages that facilitators
send in one click.

## What You Found

Student intake: {students_processed} students processed in this run.
Risk distribution at runtime:

- CRITICAL: {n_critical}
- HIGH: {n_high}
- MEDIUM: {n_medium}
- LOW: {n_low}

Duplicate student IDs removed during ingestion: {dup_count}. All data quality
issues (missing metrics, type-mismatch strings, blank notes) were auto-resolved
by the ingestion layer — no manual cleanup required.

## What You Built

A five-stage pipeline: ingest → score → LLM enrich → outputs → docs.

LLM usage this run: {api_calls} API calls, {tokens_total} tokens total
({tokens_in} input + {tokens_out} output). Fallbacks triggered: {fallbacks}.

Output files produced per run:

- intervention_priority_list.xlsx — ranked student list, colour-coded by risk
- campus_dashboard_<id>.xlsx — one tab per campus with LLM summaries
- whatsapp_messages.csv — ready-to-send parent messages (UTF-8 BOM for Arabic)
- dashboard.html — self-contained, filterable HTML dashboard (no server needed)
- word_report.docx — executive summary with per-campus tables
- docs/ — nine technical reference documents (this suite)

## What You Cut

- No ML model — would require labeled historical outcome data; deterministic
  weighted scoring is auditable and needs no training data.
- No real-time server — on-demand script run fits current facilitator workflow;
  scheduled trigger can be added later with cron or Task Scheduler.
- No OAuth or SSO — out of scope for v1; outputs are file-based, not a web app.
- No Docker or Kubernetes — single-machine Python script; no orchestration layer needed.
- No direct WhatsApp API sending — WhatsApp Business API requires business
  verification; facilitators copy-paste from CSV as the safe v1 path.

## What Next

1. Hook up real student data — replace synthetic CSVs with live exports.
2. Validate risk weights with the academic director — defaults (attendance 35%,
   practice 30%, trend 20%, notes 15%) are reasonable starting points only.
3. Confirm Arabic dialect per campus — Modern Standard vs. Gulf dialect affects
   message naturalness; update the LLM prompt system message accordingly.
4. Test outputs on LibreOffice — facilitator PCs may not have Excel; verify
   .xlsx colour-coding and column widths render correctly.
5. Consider a scheduled daily trigger — a cron job or Windows Task Scheduler
   entry replaces the manual `python main.py` step.
"""

    path.write_text(content, encoding="utf-8")
    logger.info("Wrote analysis.md: %s", path)
    return path


def _write_analysis_docx(run_log: dict, docs_dir: Path, content: dict) -> Path:
    """Write docs/analysis.docx — Word version of the 5-section analysis.

    Produces a python-docx Document with a Title heading and five level=1
    section headings mirroring analysis.md: Diagnosis, What You Found,
    What You Built, What You Cut, What Next. Driven entirely by run_log;
    the content parameter is unused (analysis is runtime-driven, not YAML-static).

    D-10 constraints: add_heading(level=0/1), add_paragraph() only — no
    OxmlElement, no custom styles, Table Grid only if tables are added.
    doc.save(str(path)) — str() required for python-docx 1.1.2 on Windows.

    Security: No student names, parent phones, or API key embedded — only
    aggregate counts from run_log (T-06-04). All run_log keys accessed via
    .get() with defaults to tolerate partial run_logs (T-06-05).

    Args:
        run_log: Pipeline run metadata dict with live numbers.
        docs_dir: Directory to write analysis.docx into.
        content: Unused — analysis.docx is fully run_log-driven, not YAML-static.

    Returns:
        Path to docs/analysis.docx.
    """
    path = docs_dir / "analysis.docx"

    # Derive stats from run_log
    tokens_used = run_log.get("tokens_used", {})
    tokens_in = tokens_used.get("input", 0)
    tokens_out = tokens_used.get("output", 0)
    tokens_total = tokens_in + tokens_out
    students_processed = run_log.get("students_processed", 0)
    api_calls = run_log.get("api_calls_made", 0)
    fallbacks = run_log.get("fallbacks_triggered", 0)
    run_timestamp = run_log.get("run_timestamp", "N/A")

    doc = Document()
    doc.add_heading("Analysis: boon-academy-intervention", level=0)
    doc.add_paragraph(f"Generated: {run_timestamp}")

    # Section 1: Diagnosis
    doc.add_heading("Diagnosis", level=1)
    doc.add_paragraph(
        "Boon Academy runs 20 campuses serving roughly 300 students. Facilitator "
        "intervention rate is currently ~30% against an 80%+ target. The gap exists "
        "because facilitators lack a fast, prioritised view of which students to contact "
        "and what to say. This pipeline closes that gap by scoring every student daily, "
        "ranking them by risk, and generating draft WhatsApp messages that facilitators "
        "send in one click."
    )

    # Section 2: What You Found
    doc.add_heading("What You Found", level=1)
    doc.add_paragraph(
        f"Student intake: {students_processed} students processed in this run. "
        "Risk distribution is available in the intervention_priority_list.xlsx output "
        "(colour-coded by CRITICAL / HIGH / MEDIUM / LOW). All data quality issues "
        "(missing metrics, type-mismatch strings, blank notes) were auto-resolved by "
        "the ingestion layer — no manual cleanup required."
    )

    # Section 3: What You Built
    doc.add_heading("What You Built", level=1)
    doc.add_paragraph(
        "A five-stage pipeline: ingest → score → LLM enrich → outputs → docs."
    )
    doc.add_paragraph(
        f"LLM usage this run: {api_calls} API calls, {tokens_total} tokens total "
        f"({tokens_in} input + {tokens_out} output). Fallbacks triggered: {fallbacks}."
    )
    doc.add_paragraph(
        "Output files: intervention_priority_list.xlsx, campus_dashboard_<id>.xlsx, "
        "whatsapp_messages.csv, dashboard.html, word_report.docx, docs/ (nine technical docs)."
    )

    # Section 4: What You Cut
    doc.add_heading("What You Cut", level=1)
    doc.add_paragraph(
        "No ML model — deterministic weighted scoring is auditable with no training data needed. "
        "No real-time server — on-demand script run fits current workflow. "
        "No OAuth or SSO — outputs are file-based, not a web app. "
        "No Docker or Kubernetes — single-machine Python script. "
        "No direct WhatsApp API sending — facilitators copy-paste from CSV as the safe v1 path."
    )

    # Section 5: What Next
    doc.add_heading("What Next", level=1)
    doc.add_paragraph(
        "1. Hook up real student data — replace synthetic CSVs with live exports.\n"
        "2. Validate risk weights with the academic director.\n"
        "3. Confirm Arabic dialect per campus (Modern Standard vs. Gulf).\n"
        "4. Test outputs on LibreOffice — facilitator PCs may not have Excel.\n"
        "5. Consider a scheduled daily trigger (cron or Windows Task Scheduler)."
    )

    # str() required for python-docx 1.1.2 on Windows (CLAUDE.md critical pitfall)
    doc.save(str(path))
    logger.info("Wrote analysis.docx: %s", path)
    return path


def _write_architecture(docs_dir: Path, content: dict) -> Path:
    """Write docs/architecture.docx — system architecture reference document.

    Hybrid structure (D-14): each section opens with a paragraph, then uses
    bullet points. Rendered via _render_doc_from_content() shared helper which
    maps type:paragraph → add_paragraph() and type:bullet → add_paragraph with
    style='List Bullet'.

    No OxmlElement. No custom styles. Table Grid only. per D-10.

    Args:
        docs_dir: Directory to write architecture.docx into.
        content: Parsed YAML dict for the 'architecture' doc block.

    Returns:
        Path to docs/architecture.docx.
    """
    path = _render_doc_from_content(content, docs_dir, "architecture.docx")
    logger.info("Wrote architecture.docx: %s", path)
    return path


def _write_security(docs_dir: Path, content: dict) -> Path:
    """Write docs/security.docx from docs_content.yaml.

    Covers env-var API key management, PII masking, data retention,
    access control. Narrative prose (D-14). No OxmlElement (D-10).

    doc.save(str(path)) used for python-docx 1.1.2 Windows compatibility
    (CLAUDE.md critical pitfall). Delegates entirely to
    _render_doc_from_content() which handles paragraph/bullet dispatch.

    Args:
        docs_dir: Directory to write security.docx into.
        content: Parsed YAML dict for the 'security' doc block.

    Returns:
        Path to docs/security.docx.
    """
    path = _render_doc_from_content(content, docs_dir, "security.docx")
    logger.info("Wrote security.docx: %s", path)
    return path


def _write_engineering_decisions(docs_dir: Path, content: dict) -> Path:
    """Write docs/engineering_decisions.docx from docs_content.yaml.

    Covers risk scoring formula rationale, LLM batching, fallback logic,
    output format choices, intentional simplicity. Narrative prose (D-14).
    No OxmlElement (D-10).

    doc.save(str(path)) used for python-docx 1.1.2 Windows compatibility
    (CLAUDE.md critical pitfall). Delegates entirely to
    _render_doc_from_content() which handles paragraph/bullet dispatch.

    Args:
        docs_dir: Directory to write engineering_decisions.docx into.
        content: Parsed YAML dict for the 'engineering_decisions' doc block.

    Returns:
        Path to docs/engineering_decisions.docx.
    """
    path = _render_doc_from_content(content, docs_dir, "engineering_decisions.docx")
    logger.info("Wrote engineering_decisions.docx: %s", path)
    return path


def _write_data_handling(docs_dir: Path, content: dict) -> Path:
    """Write docs/data_handling.docx — data handling reference document.

    Hybrid structure (D-14): each section opens with a paragraph, then uses
    bullet points. Rendered via _render_doc_from_content() shared helper which
    maps type:paragraph → add_paragraph() and type:bullet → add_paragraph with
    style='List Bullet'.

    Sections covered: Input Schema, Data Cleaning Pipeline, Missing Data Strategy,
    Edge Cases and Quality Issues, Merge Logic.

    No OxmlElement. No custom styles. Table Grid only. per D-10.

    Args:
        docs_dir: Directory to write data_handling.docx into.
        content: Parsed YAML dict for the 'data_handling' doc block.

    Returns:
        Path to docs/data_handling.docx.
    """
    path = _render_doc_from_content(content, docs_dir, "data_handling.docx")
    logger.info("Wrote data_handling.docx: %s", path)
    return path


def _write_scalability(docs_dir: Path, content: dict) -> Path:
    """Write docs/scalability.docx — scalability analysis with cost projection tables.

    Hybrid structure (D-14): each YAML section opens with a paragraph, then uses
    bullet points. After all YAML sections are rendered, a programmatic scale
    comparison table and cost projection table are appended (DOCS-07).

    Cost projection numbers are hardcoded from the demo run (D-13):
      - 31,771 tokens / 300 students = 106 tokens/student
      - 100 campuses × 15 × 40% at-risk = 600 at-risk students/run
      - 600 × 106 = 63,600 tokens/run
    Both scenarios are well within the $200/month budget target.

    No OxmlElement. No custom styles. Table Grid only. per D-10.

    Args:
        docs_dir: Directory to write scalability.docx into.
        content: Parsed YAML dict for the 'scalability' doc block.

    Returns:
        Path to docs/scalability.docx.
    """
    path = docs_dir / "scalability.docx"
    doc = Document()

    # Render title
    title = content.get("title", "Scalability Analysis")
    doc.add_heading(title, level=0)

    # Render all YAML sections using hybrid pattern (paragraph opener + bullets)
    for section in content.get("sections", []):
        heading = section.get("heading", "")
        if heading:
            doc.add_heading(heading, level=1)
        for item in section.get("body", []):
            item_type = item.get("type", "paragraph")
            text = item.get("text", "")
            if item_type == "bullet":
                doc.add_paragraph(text, style="List Bullet")
            else:
                doc.add_paragraph(text)

    # Scale comparison table (DOCS-07) — programmatic, not in YAML
    doc.add_heading("Scale Comparison", level=1)
    doc.add_paragraph(
        "The table below projects at-risk student count and token usage as campus count grows. "
        "Calculation basis: 106 tokens/student (31,771 tokens ÷ 300 students from demo run); "
        "40% CRITICAL/HIGH rate assumed."
    )
    scale_table = doc.add_table(rows=3, cols=3, style="Table Grid")
    scale_table.cell(0, 0).text = "Scale"
    scale_table.cell(0, 1).text = "At-Risk Students/Run"
    scale_table.cell(0, 2).text = "Est. Tokens/Run"
    scale_table.cell(1, 0).text = "20 campuses (current)"
    scale_table.cell(1, 1).text = "~120"
    scale_table.cell(1, 2).text = "~12,720"
    scale_table.cell(2, 0).text = "100 campuses"
    scale_table.cell(2, 1).text = "~600"
    scale_table.cell(2, 2).text = "~63,600"

    # Cost projection table
    doc.add_heading("Monthly Cost Estimate", level=1)
    doc.add_paragraph(
        "Estimated monthly cost at claude-sonnet-4-5 pricing (~$6/million tokens blended). "
        "Both scenarios are well within the $200/month budget target — 17x headroom at 100 campuses."
    )
    cost_table = doc.add_table(rows=3, cols=3, style="Table Grid")
    cost_table.cell(0, 0).text = "Runs/Month"
    cost_table.cell(0, 1).text = "Tokens/Run"
    cost_table.cell(0, 2).text = "Est. Monthly Cost"
    cost_table.cell(1, 0).text = "30"
    cost_table.cell(1, 1).text = "~12,720"
    cost_table.cell(1, 2).text = "<$1/month"
    cost_table.cell(2, 0).text = "30"
    cost_table.cell(2, 1).text = "~63,600"
    cost_table.cell(2, 2).text = "~$5-10/month"

    # str() required for python-docx 1.1.2 on Windows (CLAUDE.md critical pitfall)
    doc.save(str(path))
    logger.info("Wrote scalability.docx: %s", path)
    return path


def _write_system_design(docs_dir: Path, content: dict) -> Path:
    """Write docs/system_design.docx from docs_content.yaml.

    Covers AI role boundaries, what AI does not do, accuracy/cost/latency
    tradeoffs, human review loop recommendation, failure modes. Narrative
    prose (D-14). No OxmlElement (D-10).

    doc.save(str(path)) used for python-docx 1.1.2 Windows compatibility
    (CLAUDE.md critical pitfall). Delegates entirely to
    _render_doc_from_content() which handles paragraph/bullet dispatch.

    Args:
        docs_dir: Directory to write system_design.docx into.
        content: Parsed YAML dict for the 'system_design' doc block.

    Returns:
        Path to docs/system_design.docx.
    """
    path = _render_doc_from_content(content, docs_dir, "system_design.docx")
    logger.info("Wrote system_design.docx: %s", path)
    return path


def _write_alternatives(docs_dir: Path, content: dict) -> Path:
    """Write docs/alternatives.docx from docs_content.yaml.

    Covers risk scoring alternatives, delivery method alternatives,
    infrastructure alternatives, and what is worth building next.
    Narrative prose (D-14). No OxmlElement (D-10).

    doc.save(str(path)) used for python-docx 1.1.2 Windows compatibility
    (CLAUDE.md critical pitfall). Delegates entirely to
    _render_doc_from_content() which handles paragraph/bullet dispatch.

    Args:
        docs_dir: Directory to write alternatives.docx into.
        content: Parsed YAML dict for the 'alternatives' doc block.

    Returns:
        Path to docs/alternatives.docx.
    """
    path = _render_doc_from_content(content, docs_dir, "alternatives.docx")
    logger.info("Wrote alternatives.docx: %s", path)
    return path
